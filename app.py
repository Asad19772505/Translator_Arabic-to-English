import os
import io
import re
import zipfile
import tempfile
from pathlib import Path
from typing import List, Tuple, Dict

import streamlit as st
from dotenv import load_dotenv
from audio_recorder_streamlit import audio_recorder
from groq import Groq
import httpx

# Document handling
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image

# Optional OCR
try:
    import pytesseract
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

# -----------------------------
# Setup
# -----------------------------
st.set_page_config(page_title="Arabic ↔ English Translator (Groq)", layout="wide")
st.title("🗣️/📄 Arabic → English Translator (Groq)")

load_dotenv()
API_KEY = os.getenv("GROQ_API_KEY")

with st.sidebar:
    st.header("Settings")
    timeout_sec = st.slider("Request timeout (seconds)", 30, 6000, 180, step=30)
    st.caption("Increase if you process longer inputs.")

    model_name = st.selectbox(
        "Groq model for text translation",
        [
            "llama-3.1-70b-versatile",
            "llama-3.1-8b-instant",
            "mixtral-8x7b-32768",
        ],
        index=0,
        help="Used for document/text translation (not audio)."
    )

    polishing = st.checkbox(
        "Polish English tone (formal, concise, CFO-ready)",
        value=True,
        help="Faithful meaning, clearer English."
    )
    keep_layout = st.checkbox(
        "Preserve light structure (paragraphs & bullets)",
        value=True,
        help="Mild structure; not full formatting."
    )

    if not API_KEY:
        st.warning("Add your GROQ_API_KEY to .env or Streamlit Secrets.", icon="⚠️")

client = Groq(
    api_key=API_KEY,
    timeout=httpx.Timeout(timeout_sec, read=timeout_sec, write=timeout_sec, connect=30.0),
)

st.write("Choose **Audio** or **Document**. New: **Glossary lock**, **Batch translate**, and **Side-by-side preview**.")

top_mode = st.radio(
    "What do you want to translate?",
    ["Audio (Speech → Text/Translation)", "Document (PDF / Word / Scanned)"],
    horizontal=True,
)

# =========================
# Helpers
# =========================

def clean_whitespace(s: str) -> str:
    return re.sub(r'\s+\n', '\n', re.sub(r'[ \t]+', ' ', s)).strip()

def chunk_text(text: str, max_chars: int = 6000) -> List[str]:
    paras = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
    chunks, buf = [], ""
    for p in paras:
        if len(buf) + len(p) + 2 <= max_chars:
            buf = (buf + "\n\n" + p).strip()
        else:
            if buf:
                chunks.append(buf)
            buf = p
    if buf:
        chunks.append(buf)
    return chunks

# ---------- Glossary lock ----------
TOKEN_PREFIX = "⟦GL_"
TOKEN_SUFFIX = "⟧"

def build_glossary_list(glossary_csv: str) -> List[str]:
    # Accept comma or newline separated terms; preserve exact spacing
    raw = [t.strip() for t in re.split(r'[,\n]', glossary_csv or "") if t.strip()]
    # deduplicate, keep order
    seen, out = set(), []
    for t in raw:
        if t not in seen:
            out.append(t)
            seen.add(t)
    return out

def protect_glossary_terms(text: str, terms: List[str]) -> Tuple[str, Dict[str, str]]:
    """
    Replace exact occurrences of terms with stable tokens so the model can't alter them.
    Returns (protected_text, token_map).
    """
    token_map: Dict[str, str] = {}
    protected = text
    # Sort longer terms first to avoid partial overlaps
    for i, term in enumerate(sorted(terms, key=len, reverse=True)):
        # Use literal match (Arabic strings), word-boundary-ish but allow punctuation around
        # We'll replace all occurrences
        token = f"{TOKEN_PREFIX}{i}{TOKEN_SUFFIX}"
        token_map[token] = term
        # Use regex escaping for term
        pattern = re.escape(term)
        protected = re.sub(pattern, token, protected)
    return protected, token_map

def unprotect_tokens(text: str, token_map: Dict[str, str]) -> str:
    out = text
    for token, term in token_map.items():
        out = out.replace(token, term)
    return out

def translate_with_groq(text: str, model: str, polishing: bool, keep_layout: bool, glossary_terms: List[str]) -> str:
    if not text.strip():
        return ""

    style_instructions = [
        "Translate from Arabic to English faithfully.",
        "Do NOT summarize or omit content. Keep all data, numbers, names, and legal wording intact.",
        "Output English only, without translator notes."
    ]
    if polishing:
        style_instructions.append("Polish English for clarity, brevity, and professional (CFO) tone.")
    if keep_layout:
        style_instructions.append("Preserve paragraph breaks; convert obvious lists to bullets if appropriate.")

    glossary_instruction = ""
    if glossary_terms:
        # Explain tokens and locking
        glossary_instruction = (
            " Important: Any tokens like "
            f"{TOKEN_PREFIX}N{TOKEN_SUFFIX} are immutable placeholders for glossary terms. "
            "Never translate, change, or remove them. Output them exactly as-is."
        )

    system_prompt = (
        "You are a professional Arabic→English translator for legal/financial/business documents. "
        + " ".join(style_instructions)
        + glossary_instruction
    )

    # Protect glossary terms before sending to the model
    protected_text, token_map = protect_glossary_terms(text, glossary_terms)

    chunks = chunk_text(protected_text, max_chars=6000)
    outputs = []
    progress = st.progress(0, text="Translating…")
    for i, ch in enumerate(chunks, start=1):
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": ch}
            ],
            temperature=0.2,
        )
        outputs.append(resp.choices[0].message.content.strip())
        progress.progress(i/len(chunks), text=f"Translating chunk {i}/{len(chunks)}")
    progress.empty()

    protected_english = "\n\n".join(outputs).strip()
    # Restore glossary terms verbatim (Arabic) into the English output
    english_text = unprotect_tokens(protected_english, token_map)
    return english_text

def write_temp_file(data: bytes, suffix: str = ".wav") -> Path:
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.write(data)
    tmp.flush()
    tmp.close()
    return Path(tmp.name)

# ---------- Extraction ----------
def extract_text_from_pdf(stream: bytes, use_ocr: bool) -> Tuple[str, List[Image.Image]]:
    text_parts: List[str] = []
    ocr_images: List[Image.Image] = []
    with fitz.open(stream=stream, filetype="pdf") as doc:
        for page in doc:
            txt = page.get_text("text") or ""
            txt_clean = clean_whitespace(txt)
            if txt_clean:
                text_parts.append(txt_clean)
            elif use_ocr:
                pix = page.get_pixmap(dpi=200, alpha=False)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_images.append(img)
    joined = "\n\n".join(text_parts).strip()
    return joined, ocr_images

def extract_text_from_docx(stream: bytes) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(stream)
        tmp.flush()
        tmp_path = tmp.name
    try:
        doc = DocxDocument(tmp_path)
        paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n".join(paras)
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

def ocr_images_to_text(images: List[Image.Image], lang: str = "ara") -> str:
    if not OCR_AVAILABLE:
        st.warning("OCR libraries not available. Install pytesseract & tesseract (with Arabic) to enable OCR.", icon="⚠️")
        return ""
    texts = []
    for i, im in enumerate(images, start=1):
        st.write(f"OCR page {i}…")
        txt = pytesseract.image_to_string(im, lang=lang)
        texts.append(clean_whitespace(txt))
    return "\n\n".join([t for t in texts if t])

def extract_text_from_image(stream: bytes, use_ocr: bool) -> str:
    img = Image.open(io.BytesIO(stream)).convert("RGB")
    if use_ocr:
        return ocr_images_to_text([img], lang="ara")
    else:
        st.info("Image provided. Enable OCR to read Arabic text from image scans.")
        return ""

# ---------- Side-by-side preview ----------
def side_by_side_html(ar_text: str, en_text: str) -> str:
    # Split to paragraphs by blank lines (fallback to single newlines)
    ar_paras = [p.strip() for p in re.split(r'\n\s*\n', ar_text) if p.strip()]
    if not ar_paras:
        ar_paras = [p.strip() for p in ar_text.split("\n") if p.strip()]
    en_paras = [p.strip() for p in re.split(r'\n\s*\n', en_text) if p.strip()]
    if not en_paras:
        en_paras = [p.strip() for p in en_text.split("\n") if p.strip()]

    # Align length by padding
    n = max(len(ar_paras), len(en_paras))
    ar_paras += [""] * (n - len(ar_paras))
    en_paras += [""] * (n - len(en_paras))

    rows = []
    for i in range(n):
        rows.append(f"""
          <div class="row">
            <div class="cell ar">{ar_paras[i]}</div>
            <div class="cell en">{en_paras[i]}</div>
          </div>
        """)

    html = f"""
    <style>
      .grid {{
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 12px;
        font-family: ui-sans-serif, system-ui, -apple-system, "Segoe UI", Roboto, "Helvetica Neue", Arial;
      }}
      .hdr {{
        font-weight: 600; margin-bottom: 8px; font-size: 14px;
      }}
      .row {{
        display: contents;
      }}
      .cell {{
        border: 1px solid #e5e7eb;
        border-radius: 8px;
        padding: 12px;
        background: #fff;
        white-space: pre-wrap;
        line-height: 1.5;
      }}
      .ar {{ direction: rtl; text-align: right; }}
      .en {{ direction: ltr; text-align: left; }}
    </style>
    <div class="grid">
      <div class="hdr">Arabic (source)</div>
      <div class="hdr">English (translation)</div>
      {''.join(rows)}
    </div>
    """
    return html

# =========================
# AUDIO (unchanged)
# =========================
if top_mode == "Audio (Speech → Text/Translation)":
    mode = st.radio(
        "Audio Mode",
        ["Translate Arabic → English", "Transcribe Arabic (Arabic text)"],
        horizontal=True,
    )

    col1, col2 = st.columns(2, vertical_alignment="center")
    with col1:
        st.subheader("🎧 Microphone (browser)")
        st.caption("Click once to start, click again to stop. Auto-stops after silence.")
        audio_bytes = audio_recorder(pause_threshold=2.0)

    with col2:
        st.subheader("📤 Or upload audio")
        uploaded_audio = st.file_uploader(
            "Supported: wav, mp3, m4a, ogg, webm", type=["wav", "mp3", "m4a", "ogg", "webm"]
        )

    prompt = st.text_input(
        "Optional context (names, terms, spellings)",
        help="Helps the model with unusual words or spellings.",
    )

    go = st.button("Process Audio", type="primary", use_container_width=True)

    def process_with_groq_audio(audio_path: Path, mode: str, prompt: str | None = None) -> str:
        if mode.startswith("Translate"):
            with open(audio_path, "rb") as f:
                res = client.audio.translations.create(
                    file=(audio_path.name, f.read()),
                    model="whisper-large-v3",
                    response_format="json",
                    temperature=0.0,
                    prompt=prompt or None,
                )
            return res.text

        with open(audio_path, "rb") as f:
            res = client.audio.transcriptions.create(
                file=f,
                model="whisper-large-v3-turbo",
                response_format="json",
                temperature=0.0,
                language="ar",
                prompt=prompt or None,
            )
        return res.text

    if go:
        if not API_KEY:
            st.error("GROQ_API_KEY not found. Set it in .env or Streamlit Secrets.", icon="🚫")
            st.stop()

        temp_path = None
        try:
            if audio_bytes:
                temp_path = write_temp_file(audio_bytes, suffix=".wav")
            elif uploaded_audio is not None:
                suffix = "." + uploaded_audio.name.split(".")[-1].lower()
                temp_path = write_temp_file(uploaded_audio.getbuffer(), suffix=suffix)
            else:
                st.warning("Please record or upload an audio file.", icon="ℹ️")
                st.stop()

            with st.spinner("Transcribing/Translating with Groq Whisper…"):
                text = process_with_groq_audio(temp_path, mode, prompt)

            st.success("Done.")
            st.subheader("Result")
            st.text_area("Output", value=text, height=220)

            st.download_button(
                "⬇️ Download as .txt",
                data=text.encode("utf-8"),
                file_name="transcript.txt",
                mime="text/plain",
                use_container_width=True,
            )

        except httpx.ReadTimeout:
            st.error("Request timed out. Increase the timeout slider and try again.")
        except Exception as e:
            st.exception(e)
        finally:
            if temp_path and temp_path.exists():
                try:
                    temp_path.unlink()
                except Exception:
                    pass

# =========================
# DOCUMENTS (single + batch)
# =========================
else:
    st.subheader("📄 Upload Document(s) (Arabic → English)")
    st.caption("Supported: PDF, DOCX, Images (PNG/JPG/JPEG/TIFF). Toggle OCR for scans.")

    tab_single, tab_batch = st.tabs(["Single file", "Batch translate (multi-upload)"])

    # ---- Shared options
    use_ocr = st.toggle(
        "Enable OCR (for scanned PDFs/images)", value=False,
        help="Requires Tesseract + Arabic language data on the server."
    )
    glossary_csv = st.text_area(
        "Glossary lock (Arabic terms to preserve verbatim)",
        help="Comma or newline separated. These Arabic strings will appear unchanged in the English output.",
        placeholder="e.g., شركة الأقطار, وزارة المالية, نيوم"
    )
    glossary_terms = build_glossary_list(glossary_csv)

    # ===== Single file =====
    with tab_single:
        file = st.file_uploader(
            "Upload a file", type=["pdf", "docx", "png", "jpg", "jpeg", "tif", "tiff", "bmp"], key="single_up"
        )
        show_preview = st.checkbox("Show side-by-side HTML preview after translation", value=True, key="single_preview")
        want_bilingual_docx = st.checkbox("Generate bilingual DOCX (source + translation)", value=False, key="single_docx")

        go_doc = st.button("Translate Document", type="primary", use_container_width=True, key="single_go")

        if go_doc:
            if not API_KEY:
                st.error("GROQ_API_KEY not found. Set it in .env or Streamlit Secrets.", icon="🚫")
                st.stop()
            if not file:
                st.warning("Please upload a document.", icon="ℹ️")
                st.stop()

            name = file.name.lower()
            data = file.getbuffer()

            try:
                with st.spinner("Extracting text…"):
                    arabic_text = ""
                    if name.endswith(".pdf"):
                        pdf_text, imgs_for_ocr = extract_text_from_pdf(data, use_ocr=use_ocr)
                        arabic_text = pdf_text
                        if not arabic_text and use_ocr and imgs_for_ocr:
                            arabic_text = ocr_images_to_text(imgs_for_ocr, lang="ara")
                    elif name.endswith(".docx"):
                        arabic_text = extract_text_from_docx(data)
                    else:
                        arabic_text = extract_text_from_image(data, use_ocr=use_ocr)

                    arabic_text = clean_whitespace(arabic_text)

                if not arabic_text:
                    st.error("No text found. If this is a scan or image-based PDF, enable OCR.")
                    st.stop()

                with st.spinner("Translating with Groq…"):
                    english_text = translate_with_groq(
                        arabic_text, model=model_name,
                        polishing=polishing, keep_layout=keep_layout,
                        glossary_terms=glossary_terms
                    )

                st.success("Translation complete.")
                colL, colR = st.columns([1,1])
                with colL:
                    st.subheader("English Translation")
                    st.text_area("Output", value=english_text, height=360, key="single_out")
                    st.download_button(
                        "⬇️ Download translation (.txt)",
                        data=english_text.encode("utf-8"),
                        file_name=f"{Path(file.name).stem}_EN.txt",
                        mime="text/plain",
                        use_container_width=True,
                    )
                with colR:
                    if show_preview:
                        st.subheader("Side-by-side preview")
                        html = side_by_side_html(arabic_text, english_text)
                        st.components.v1.html(html, height=520, scrolling=True)

                if want_bilingual_docx:
                    docx_buf = io.BytesIO()
                    doc = DocxDocument()
                    doc.add_heading(f"Translation of: {file.name}", level=1)
                    doc.add_paragraph("Source Language: Arabic")
                    doc.add_paragraph("Target Language: English")
                    doc.add_paragraph("-" * 30)
                    doc.add_heading("English Translation", level=2)
                    for para in english_text.split("\n"):
                        doc.add_paragraph(para)
                    doc.save(docx_buf)
                    st.download_button(
                        "⬇️ Download bilingual DOCX",
                        data=docx_buf.getvalue(),
                        file_name=f"{Path(file.name).stem}_EN.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

            except httpx.ReadTimeout:
                st.error("Request timed out. Increase the timeout slider and try again.")
            except Exception as e:
                st.exception(e)

    # ===== Batch translate =====
    with tab_batch:
        files = st.file_uploader(
            "Upload multiple files", accept_multiple_files=True,
            type=["pdf", "docx", "png", "jpg", "jpeg", "tif", "tiff", "bmp"],
            key="batch_up"
        )
        gen_bilingual = st.checkbox("Also generate bilingual DOCX for each file", value=False, key="batch_docx")
        show_batch_preview = st.checkbox("Show a quick side-by-side preview for the first file", value=True, key="batch_preview")

        go_batch = st.button("Translate All", type="primary", use_container_width=True, key="batch_go")

        if go_batch:
            if not API_KEY:
                st.error("GROQ_API_KEY not found. Set it in .env or Streamlit Secrets.", icon="🚫")
                st.stop()
            if not files:
                st.warning("Please upload at least one file.", icon="ℹ️")
                st.stop()

            zip_buf = io.BytesIO()
            previews_done = False
            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for idx, file in enumerate(files, start=1):
                    name = file.name
                    lname = name.lower()
                    data = file.getbuffer()
                    st.write(f"Processing **{name}** ({idx}/{len(files)}) …")

                    try:
                        # Extract
                        if lname.endswith(".pdf"):
                            pdf_text, imgs_for_ocr = extract_text_from_pdf(data, use_ocr=use_ocr)
                            arabic_text = pdf_text or ""
                            if not arabic_text and use_ocr and imgs_for_ocr:
                                arabic_text = ocr_images_to_text(imgs_for_ocr, lang="ara")
                        elif lname.endswith(".docx"):
                            arabic_text = extract_text_from_docx(data)
                        else:
                            arabic_text = extract_text_from_image(data, use_ocr=use_ocr)
                        arabic_text = clean_whitespace(arabic_text)

                        if not arabic_text:
                            st.warning(f"No text found in {name}. Skipping.", icon="⚠️")
                            continue

                        # Translate
                        english_text = translate_with_groq(
                            arabic_text, model=model_name,
                            polishing=polishing, keep_layout=keep_layout,
                            glossary_terms=glossary_terms
                        )

                        # Add TXT to ZIP
                        zf.writestr(f"{Path(name).stem}_EN.txt", english_text)

                        # Optional bilingual DOCX
                        if gen_bilingual:
                            docx_io = io.BytesIO()
                            doc = DocxDocument()
                            doc.add_heading(f"Translation of: {name}", level=1)
                            doc.add_paragraph("Source Language: Arabic")
                            doc.add_paragraph("Target Language: English")
                            doc.add_paragraph("-" * 30)
                            doc.add_heading("English Translation", level=2)
                            for para in english_text.split("\n"):
                                doc.add_paragraph(para)
                            doc.save(docx_io)
                            zf.writestr(f"{Path(name).stem}_EN.docx", docx_io.getvalue())

                        # Optional quick preview for first file
                        if show_batch_preview and not previews_done:
                            st.subheader("Preview (first translated file)")
                            html = side_by_side_html(arabic_text, english_text)
                            st.components.v1.html(html, height=520, scrolling=True)
                            previews_done = True

                    except httpx.ReadTimeout:
                        st.error(f"Timeout translating {name}. Consider increasing the timeout.", icon="⏱️")
                    except Exception as e:
                        st.exception(e)

            st.success("Batch translation complete.")
            st.download_button(
                "⬇️ Download all translations (ZIP)",
                data=zip_buf.getvalue(),
                file_name="translations.zip",
                mime="application/zip",
                use_container_width=True,
            )

# -----------------------------
# Tips
# -----------------------------
st.markdown(
    """
---
**Tips**
- For scanned PDFs/images, enable **OCR** (requires `tesseract-ocr` + `tesseract-ocr-ara` and Python `pytesseract`, `Pillow`).
- The **Glossary lock** protects exact Arabic strings using internal placeholders. They will appear **verbatim** in English output.
- Very long documents are chunked automatically before translation.
"""
)
